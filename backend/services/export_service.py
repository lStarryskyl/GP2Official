"""Export service for PDF and DOCX generation."""

from io import BytesIO
from typing import List, Dict, Any
from datetime import datetime

from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.colors import HexColor

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


class ExportService:
    """Service for exporting projects to PDF and DOCX."""
    
    # Acorn brand colors
    ORANGE = HexColor('#F5A623')
    BLUE = HexColor('#4A7BA7')
    NAVY = HexColor('#1B2D45')
    
    async def export_project_pdf(
        self,
        project: Dict[str, Any],
        requirements: List[Dict[str, Any]],
        tasks: List[Dict[str, Any]]
    ) -> BytesIO:
        """Export complete project to PDF with Acorn branding."""
        
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.5*inch)
        story = []
        styles = getSampleStyleSheet()
        
        # Custom styles with Acorn colors
        title_style = ParagraphStyle(
            'AcornTitle',
            parent=styles['Heading1'],
            fontSize=28,
            textColor=self.NAVY,
            spaceAfter=20,
            fontName='Helvetica-Bold'
        )
        
        heading_style = ParagraphStyle(
            'AcornHeading',
            parent=styles['Heading2'],
            fontSize=18,
            textColor=self.ORANGE,
            spaceAfter=12,
            spaceBefore=12,
            fontName='Helvetica-Bold'
        )
        
        # Header with Acorn branding
        header = Paragraph("Acorn - AI Planning Platform", styles['Normal'])
        story.append(header)
        story.append(Spacer(1, 0.3*inch))
        
        # Project Title
        story.append(Paragraph(project.get('name', 'Untitled Project'), title_style))
        story.append(Paragraph(project.get('description', ''), styles['Normal']))
        story.append(Spacer(1, 0.2*inch))
        
        # Project Metadata
        metadata = [
            ['Owner', project.get('owner', 'N/A')],
            ['Status', project.get('status', 'N/A')],
            ['Created', str(project.get('created_at', 'N/A'))[:10]],
            ['Phase', project.get('current_phase', 'N/A')]
        ]
        
        metadata_table = Table(metadata, colWidths=[2*inch, 4*inch])
        metadata_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), self.ORANGE),
            ('TEXTCOLOR', (0, 0), (0, -1), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('TOPPADDING', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 1, colors.grey)
        ]))
        story.append(metadata_table)
        story.append(Spacer(1, 0.3*inch))
        
        # Requirements Section
        if requirements:
            story.append(Paragraph("Requirements", heading_style))
            
            req_data = [['ID', 'Title', 'Type', 'Priority', 'Status']]
            for req in requirements[:50]:  # Limit to avoid huge PDFs
                req_data.append([
                    req.get('id', '')[:8],
                    req.get('title', 'N/A')[:40],
                    req.get('type', 'N/A'),
                    req.get('priority', 'N/A'),
                    req.get('status', 'N/A')
                ])
            
            req_table = Table(req_data, colWidths=[0.8*inch, 3*inch, 1*inch, 1*inch, 1*inch])
            req_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), self.ORANGE),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 11),
                ('FONTSIZE', (0, 1), (-1, -1), 9),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey])
            ]))
            story.append(req_table)
            story.append(Spacer(1, 0.3*inch))
        
        # Tasks Section
        if tasks:
            story.append(Paragraph("Tasks", heading_style))
            
            task_data = [['Title', 'Assignee', 'Priority', 'Status']]
            for task in tasks[:50]:
                task_data.append([
                    task.get('title', 'N/A')[:50],
                    task.get('assignee', 'Unassigned'),
                    task.get('priority', 'N/A'),
                    task.get('status', 'N/A')
                ])
            
            task_table = Table(task_data, colWidths=[3.5*inch, 1.5*inch, 1*inch, 1*inch])
            task_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), self.BLUE),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 11),
                ('FONTSIZE', (0, 1), (-1, -1), 9),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey])
            ]))
            story.append(task_table)
        
        # Footer
        story.append(Spacer(1, 0.5*inch))
        footer_text = f"Generated by Acorn on {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}"
        story.append(Paragraph(footer_text, styles['Italic']))
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        return buffer
    
    async def export_project_docx(
        self,
        project: Dict[str, Any],
        requirements: List[Dict[str, Any]],
        tasks: List[Dict[str, Any]]
    ) -> BytesIO:
        """Export complete project to DOCX with Acorn branding."""
        
        doc = Document()
        
        # Header
        header = doc.add_heading('Acorn - AI Planning Platform', level=3)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Project Title
        title = doc.add_heading(project.get('name', 'Untitled Project'), level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        title_run = title.runs[0]
        title_run.font.color.rgb = RGBColor(27, 45, 69)  # Navy
        
        # Description
        doc.add_paragraph(project.get('description', ''))
        
        # Metadata
        doc.add_heading('Project Information', level=2)
        metadata_table = doc.add_table(rows=4, cols=2)
        metadata_table.style = 'Light Grid Accent 1'
        
        metadata = [
            ('Owner', project.get('owner', 'N/A')),
            ('Status', project.get('status', 'N/A')),
            ('Created', str(project.get('created_at', 'N/A'))[:10]),
            ('Phase', project.get('current_phase', 'N/A'))
        ]
        
        for i, (key, value) in enumerate(metadata):
            metadata_table.rows[i].cells[0].text = key
            metadata_table.rows[i].cells[1].text = str(value)
        
        # Requirements
        if requirements:
            doc.add_page_break()
            req_heading = doc.add_heading('Requirements', level=2)
            req_heading.runs[0].font.color.rgb = RGBColor(245, 166, 35)  # Orange
            
            req_table = doc.add_table(rows=1, cols=5)
            req_table.style = 'Light Grid Accent 1'
            
            # Header row
            hdr_cells = req_table.rows[0].cells
            hdr_cells[0].text = 'ID'
            hdr_cells[1].text = 'Title'
            hdr_cells[2].text = 'Type'
            hdr_cells[3].text = 'Priority'
            hdr_cells[4].text = 'Status'
            
            # Data rows
            for req in requirements[:50]:
                row_cells = req_table.add_row().cells
                row_cells[0].text = req.get('id', '')[:8]
                row_cells[1].text = req.get('title', 'N/A')
                row_cells[2].text = req.get('type', 'N/A')
                row_cells[3].text = req.get('priority', 'N/A')
                row_cells[4].text = req.get('status', 'N/A')
        
        # Tasks
        if tasks:
            doc.add_page_break()
            task_heading = doc.add_heading('Tasks', level=2)
            task_heading.runs[0].font.color.rgb = RGBColor(74, 123, 167)  # Blue
            
            task_table = doc.add_table(rows=1, cols=4)
            task_table.style = 'Light Grid Accent 1'
            
            # Header row
            hdr_cells = task_table.rows[0].cells
            hdr_cells[0].text = 'Title'
            hdr_cells[1].text = 'Assignee'
            hdr_cells[2].text = 'Priority'
            hdr_cells[3].text = 'Status'
            
            # Data rows
            for task in tasks[:50]:
                row_cells = task_table.add_row().cells
                row_cells[0].text = task.get('title', 'N/A')
                row_cells[1].text = task.get('assignee', 'Unassigned')
                row_cells[2].text = task.get('priority', 'N/A')
                row_cells[3].text = task.get('status', 'N/A')
        
        # Footer
        doc.add_paragraph()
        footer = doc.add_paragraph(f"Generated by Acorn on {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}")
        footer.runs[0].font.italic = True
        footer.runs[0].font.size = Pt(9)
        
        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
